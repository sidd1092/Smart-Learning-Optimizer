# from docx import Document
# from docx.shared import Inches
# import requests
# import tkinter as tk
# from tkinter import filedialog
# import openai
# import requests

# def rephrase_text(text):
#     # Set your OpenAI API key here
#     openai.api_key = "your_openai_api_key"

#     try:
#         # Constructing the prompt to ask for a simpler rephrasing
#         prompt = f"Please rephrase the following text in simpler terms:\n\n{text}"

#         # Making the API call to OpenAI's GPT model
#         response = openai.Completion.create(
#           engine="text-davinci-003",  # You can choose a different model based on your needs
#           prompt=prompt,
#           max_tokens=100,  # Adjust based on the expected length of the rephrased text
#           temperature=0.7,  # Adjust for creativity. Lower values make the output more deterministic
#           top_p=1,
#           frequency_penalty=0,
#           presence_penalty=0
#         )

#         # Extracting the rephrased text from the response
#         rephrased_text = response.choices[0].text.strip()

#         return rephrased_text
#     except Exception as e:
#         print(f"An error occurred: {e}")
#         return text

# def process_word_document():
#     # Create a dialog box to select input file
#     input_path = filedialog.askopenfilename(title="Select Input File", filetypes=(("Word Documents", "*.docx"), ("All Files", "*.*")))
#     if not input_path:
#         return

#     # Create a dialog box to select output file
#     output_path = filedialog.asksaveasfilename(title="Save Output File", defaultextension=".docx", filetypes=(("Word Documents", "*.docx"), ("All Files", "*.*")))
#     if not output_path:
#         return

#     # Load the Word document
#     doc = Document(input_path)
#     new_doc = Document()

#     # Process each paragraph
#     for para in doc.paragraphs:
#         rephrased_text = rephrase_text(para.text)
#         new_doc.add_paragraph(rephrased_text)

#     # Process each image
#     # Note: python-docx doesn't support extracting images directly by their positions.
#     # This is a simplified approach to re-add images assuming they are separate from text.
#     # For exact positioning, more advanced handling is required.
#     for shape in doc.inline_shapes:
#         if shape.type == 3:
#             new_doc.add_picture(shape.image.filename, width=Inches(2))

#     # Save the new document
#     new_doc.save(output_path)

# # Create a dialog box to initiate the process
# root = tk.Tk()
# root.withdraw()
# process_word_document()




from docx import Document
from docx.shared import Inches
import tkinter as tk
from tkinter import filedialog
import openai

def rephrase_text_with_chat(text):
    # Set your OpenAI API key here
    openai.api_key = "sk-proj-wUyed3Ao0Fvy9WVGwwaUT3BlbkFJqUZt1Jywu1so3wT73rMl"

    try:
        # Constructing the messages for chat-based rephrasing
        messages = [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": f"Please rephrase the following text in simpler terms:\n\n{text}"}
        ]

        # Making the API call to OpenAI's GPT model using chat
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo-16k",  # You can choose a different model based on your needs
            messages=messages
        )

        # Extracting the rephrased text from the response
        rephrased_text = response.choices[0].message['content'].strip()

        return rephrased_text
    except Exception as e:
        print(f"An error occurred: {e}")
        return text

def process_word_document():
    # Create a dialog box to select input file
    input_path = filedialog.askopenfilename(title="Select Input File", filetypes=(("Word Documents", "*.docx"), ("All Files", "*.*")))
    if not input_path:
        return

    # Create a dialog box to select output file
    output_path = filedialog.asksaveasfilename(title="Save Output File", defaultextension=".docx", filetypes=(("Word Documents", "*.docx"), ("All Files", "*.*")))
    if not output_path:
        return

    # Load the Word document
    doc = Document(input_path)
    new_doc = Document()

    # Process each paragraph
    for para in doc.paragraphs:
        rephrased_text = rephrase_text_with_chat(para.text)
        new_doc.add_paragraph(rephrased_text)

    # Process each image
    # Note: python-docx doesn't support extracting images directly by their positions.
    # This is a simplified approach to re-add images assuming they are separate from text.
    # For exact positioning, more advanced handling is required.
    for shape in doc.inline_shapes:
        if shape.type == 3:
            new_doc.add_picture(shape.image.filename, width=Inches(2))

    # Save the new document
    new_doc.save(output_path)

# Create a dialog box to initiate the process
root = tk.Tk()
root.withdraw()
process_word_document()