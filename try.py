import fitz  # PyMuPDF
from docx import Document
import pytesseract
from PIL import Image
import io

# Configure pytesseract path to your installation path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def pdf_to_word(pdf_path, word_path):
    # Open the PDF
    pdf = fitz.open(pdf_path)
    doc = Document()

    for page_num in range(len(pdf)):
        page = pdf.load_page(page_num)
        
        # Extract text
        text = page.get_text("text")
        doc.add_paragraph(text)
        
        # Extract images
        image_list = page.get_images(full=True)
        for image_index, img in enumerate(page.get_images(full=True)):
            # Get the image XREF
            xref = img[0]
            base_image = pdf.extract_image(xref)
            image_bytes = base_image["image"]

            # Convert it to a PIL Image
            image = Image.open(io.BytesIO(image_bytes))
            
            # Perform OCR
            ocr_text = pytesseract.image_to_string(image)
            doc.add_paragraph(ocr_text)
    
    # Save the Word document
    doc.save(word_path)

# Example usage
pdf_to_word("input.pdf", "output.docx")