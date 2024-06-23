from PIL import Image
import io
import tempfile
import os
from docx.shared import Inches
import fitz  # PyMuPDF
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox

def convert_pdf_to_word_with_images(pdf_path, word_path):
    document = Document()
    document.add_heading('PDF to Word Conversion', 0)

    pdf = fitz.open(pdf_path)
    for page_num in range(len(pdf)):
        page = pdf.load_page(page_num)
        text = page.get_text("text")
        paragraphs = text.split('\n\n')  # Assuming double line breaks for new paragraphs

        for para in paragraphs:
            if para.startswith(('-', '•', '*')) or para[0].isdigit() and para[1] == '.':
                # This is a simplistic check for bullet points; may need refinement
                p = document.add_paragraph()
                p.add_run(para)
                p.style = 'ListBullet'
            else:
                document.add_paragraph(para)

        for image_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            try:
                base_image = pdf.extract_image(xref)
                image_bytes = base_image["image"]

                image_stream = io.BytesIO(image_bytes)
                image = Image.open(image_stream)

                dpi_x, dpi_y = image.info.get('dpi', (132, 132))
                width_inches = image.width / dpi_x
                height_inches = image.height / dpi_y

                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_image:
                    image.save(tmp_image.name)
                    document.add_picture(tmp_image.name, width=Inches(width_inches), height=Inches(height_inches))

                os.unlink(tmp_image.name)
            except Exception as e:
                print(f"Error processing image {image_index} on page {page_num}: {e}")

    document.save(word_path)

def convert_pdf_to_word_with_dialog():
    root = tk.Tk()
    root.withdraw()
    pdf_path = filedialog.askopenfilename(title="Select PDF File", filetypes=[("PDF files", "*.pdf")])
    if not pdf_path:
        messagebox.showerror("Error", "No PDF file selected!")
        return

    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if not save_path:
        messagebox.showerror("Error", "No save path specified!")
        return

    convert_pdf_to_word_with_images(pdf_path, save_path)
    messagebox.showinfo("Success", "PDF has been converted to Word successfully!")

if __name__ == "__main__":
    convert_pdf_to_word_with_dialog()